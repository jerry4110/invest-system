"""T-36 수용 기준: 이탈도(FR-05-02)·LLM 제안 근거 필수(FR-05-14)·정합성 재검증·전후 비교(FR-05-15)·리포트."""
import json

import pytest
from fastapi.testclient import TestClient

from backend import main as main_mod
from backend.infra import db as db_mod

CSV = """종목명,종목코드,카테고리,지역,보유수량,평균매입가,현재가,매입금액,평가금액,평가손익,수익률
삼성전자,005930,반도체,국내,100,70000,75000,7000000,7500000,500000,7.14
엔비디아,NVDA,반도체,해외,10,120000,180000,1200000,1800000,600000,50%
"""


class CapturingLLM:
    model = "fake"
    def __init__(self, text): self._t, self.prompts = text, []
    def complete(self, prompt, max_tokens=1024):
        from backend.adapters.llm.base import LLMResult
        self.prompts.append(prompt)
        return LLMResult(text=self._t, model="fake", input_tokens=50,
                         output_tokens=50, cost_usd=0.001)


@pytest.fixture()
def client(tmp_path, monkeypatch):
    monkeypatch.setattr(db_mod, "_engine", None)
    monkeypatch.setattr(db_mod, "_SessionLocal", None)
    db_mod.init_db(str(tmp_path / "t.db"))
    c = TestClient(main_mod.create_app())
    c.post("/api/portfolio/upload",
           files={"file": ("잔고.csv", CSV.encode("utf-8-sig"), "text/csv")})
    c.put("/api/portfolio/cash", json={"amount": 700000})   # 총자산 1,000만
    c.put("/api/strategy/allocation",
          json={"stock_pct": 70, "cash_pct": 30, "domestic_pct": 60, "overseas_pct": 40})
    return c


def test_deviation_computed(client):
    """주식 93%(목표70) → +23%p 이탈."""
    body = client.get("/api/rebalance/deviation").json()
    dev = {d["key"]: d for d in body["deviations"]}
    assert dev["stock_pct"]["current_pct"] == 93.0
    assert dev["stock_pct"]["deviation_pp"] == 23.0
    assert dev["cash_pct"]["deviation_pp"] == -23.0
    # 국내/해외는 주식 내 비중: 국내 7.5M/9.3M = 80.65%
    assert dev["domestic_pct"]["current_pct"] == pytest.approx(80.65, abs=0.01)
    assert body["as_of"]


PROPOSAL = json.dumps({
    "actions": [
        {"ticker": "005930", "name": "삼성전자", "action": "매도", "qty": 30,
         "est_amount": 2250000, "rationale": "주식 비중 초과분 축소 + 반도체 집중 완화"},
    ],
    "target_cash_pct": 30,
    "summary": "주식 93%→70%로 축소, 현금 확보",
}, ensure_ascii=False)


def test_propose_with_validation_and_history(client, monkeypatch):
    from backend.services import rebalance_service
    llm = CapturingLLM(PROPOSAL)
    monkeypatch.setattr(rebalance_service, "_default_adapter", lambda: llm)

    body = client.post("/api/rebalance/propose").json()
    assert body["actions"][0]["rationale"]                   # FR-05-14 근거 필수
    # 전/후 비교 (FR-05-15): 매도 225만 → 주식 705만/현금 295만
    after = {a["key"]: a["after_pct"] for a in body["before_after"]}
    assert after["stock_pct"] == pytest.approx(70.5, abs=0.1)
    assert body["warnings"] == []                            # 정합성 통과
    assert "참고자료" in body["disclaimer"]
    # 전략 컨텍스트 포함 + 이력 저장 (kind=rebalance)
    assert "이탈" in llm.prompts[0] or "목표" in llm.prompts[0]
    from backend.infra.schema import AnalysisResult
    with db_mod.get_session() as s:
        assert s.query(AnalysisResult).filter_by(kind="rebalance").count() == 1


def test_propose_flags_inconsistent_amounts(client, monkeypatch):
    """LLM 수량 오류 → 경고 표시 (숨기지 않음)."""
    bad = json.dumps({"actions": [
        {"ticker": "005930", "name": "삼성전자", "action": "매도", "qty": 500,
         "est_amount": 37500000, "rationale": "r"}],
        "target_cash_pct": 30, "summary": "s"}, ensure_ascii=False)
    from backend.services import rebalance_service
    monkeypatch.setattr(rebalance_service, "_default_adapter", lambda: CapturingLLM(bad))
    body = client.post("/api/rebalance/propose").json()
    assert any("보유 수량" in w or "초과" in w for w in body["warnings"])


def test_missing_rationale_flagged(client, monkeypatch):
    noR = json.dumps({"actions": [
        {"ticker": "005930", "name": "삼성전자", "action": "매도", "qty": 10,
         "est_amount": 750000, "rationale": ""}],
        "target_cash_pct": 30, "summary": "s"}, ensure_ascii=False)
    from backend.services import rebalance_service
    monkeypatch.setattr(rebalance_service, "_default_adapter", lambda: CapturingLLM(noR))
    body = client.post("/api/rebalance/propose").json()
    assert any("근거" in w for w in body["warnings"])        # FR-05-14


def test_rebalance_report_docx(client, monkeypatch, tmp_path):
    """FR-05-17: 리밸런싱 Word 리포트."""
    import docx as docxlib
    from backend.services import rebalance_service, report_service
    monkeypatch.setattr(report_service, "REPORT_DIR", tmp_path / "reports")
    monkeypatch.setattr(rebalance_service, "_default_adapter", lambda: CapturingLLM(PROPOSAL))
    client.post("/api/rebalance/propose")

    r = client.post("/api/reports/rebalance")
    assert r.status_code == 200
    path = tmp_path / "reports" / r.json()["relpath"]
    text = "\n".join(p.text for p in docxlib.Document(path).paragraphs)
    tables = "\n".join(c.text for t in docxlib.Document(path).tables
                       for row in t.rows for c in row.cells)
    assert "삼성전자" in tables and "매도" in tables
    assert "참고자료" in text
