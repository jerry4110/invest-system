"""T-27 수용 기준: docx 리포트 생성(고지문·출처 포함), 보관함 목록·다운로드."""
import json

import pytest
from fastapi.testclient import TestClient

from backend import main as main_mod
from backend.infra import db as db_mod


@pytest.fixture()
def env(tmp_path, monkeypatch):
    monkeypatch.setattr(db_mod, "_engine", None)
    monkeypatch.setattr(db_mod, "_SessionLocal", None)
    db_mod.init_db(str(tmp_path / "t.db"))
    from backend.services import report_service
    monkeypatch.setattr(report_service, "REPORT_DIR", tmp_path / "reports")
    # 분석 이력 시드 (T-26 저장 형식)
    from backend.services.analysis_service import _save_result
    _save_result("005930", "comprehensive", {
        "recommendation": "매수", "fair_value_current": 320000,
        "plan": "300,000원 이하 3회 분할", "rationale": "재무 양호",
        "assumptions": ["PER 밴드 12~18배"]})
    _save_result("005930", "debate", {"bull": "성장 지속", "bear": "밸류 부담",
                                      "conclusion": "분할 접근"})
    # 라이브 분석 로더는 모킹
    monkeypatch.setattr(report_service, "_load_fundamental", lambda t: {
        "financials": [{"year": 2025, "revenue": 333e12, "operating_profit": 43e12,
                        "net_income": 45e12}],
        "evaluation": {"items": [{"label": "ROE", "value": 12.3, "status": "미충족",
                                  "threshold": 15, "direction": "min", "metric": "roe_pct"}],
                       "tier1": {"verdict": "미충족", "passed": [], "failed": ["roe_pct"],
                                 "unknown": []}}})
    monkeypatch.setattr(report_service, "_load_technical", lambda t: {
        "signal": {"verdict": "중립", "reasons": ["혼조"]}, "rsi": 44.0,
        "ma_alignment": "혼조"})
    monkeypatch.setattr(report_service, "_load_news", lambda t: {
        "news": [{"title": "HBM4 양산", "link": "https://n/1", "source": "전자신문",
                  "date": "2026-07-10", "sentiment": "호재"}],
        "disclosures": [{"title": "자기주식취득", "link": "https://dart.fss.or.kr/x",
                         "date": "2026-07-10"}], "consensus": None})
    return tmp_path


def test_generate_docx_report(env):
    """FR-09-01·04-41: docx 생성 — 핵심 내용·출처·고지문 포함."""
    import docx
    from backend.services import report_service

    meta = report_service.generate_stock_report("005930")
    path = env / "reports" / meta["relpath"]
    assert path.exists() and path.suffix == ".docx"

    text = "\n".join(p.text for p in docx.Document(path).paragraphs)
    tables_text = "\n".join(c.text for t in docx.Document(path).tables
                            for r in t.rows for c in r.cells)
    assert "005930" in text
    assert "매수" in text                                  # 종합 판단
    assert "분할 접근" in text                             # 토론 결론
    assert "https://n/1" in text or "https://n/1" in tables_text   # 출처 (FR-04-21)
    assert "참고자료" in text                               # 고지문 (FR-04-36)
    assert "ROE" in tables_text                             # 지표 표


def test_report_registry_and_download(env):
    from backend.services import report_service
    report_service.generate_stock_report("005930")
    client = TestClient(main_mod.create_app())

    lst = client.get("/api/reports").json()
    assert len(lst) == 1 and lst[0]["ticker"] == "005930"
    rid = lst[0]["id"]
    res = client.get(f"/api/reports/{rid}/download")
    assert res.status_code == 200
    assert "wordprocessingml" in res.headers["content-type"]
    assert res.content[:2] == b"PK"                         # docx(zip) 시그니처
    assert client.get("/api/reports/9999/download").status_code == 404


def test_generate_via_api(env):
    client = TestClient(main_mod.create_app())
    r = client.post("/api/reports/stock/005930")
    assert r.status_code == 200 and r.json()["relpath"].endswith(".docx")
