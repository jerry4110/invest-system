"""리포트 생성 서비스 (M9, FR-09-01~03) — python-docx 표준 템플릿."""
import json
import logging
from datetime import date, datetime
from pathlib import Path

from backend.infra.db import get_session
from backend.infra.schema import Report

logger = logging.getLogger(__name__)
REPORT_DIR = Path("reports")

_DISCLAIMER = "※ 본 리포트는 투자 참고자료이며 최종 판단은 투자자 본인의 책임입니다."


def _load_fundamental(ticker: str) -> dict | None:
    from backend.services.analysis_service import analyze_fundamental
    return analyze_fundamental(ticker)


def _load_technical(ticker: str) -> dict | None:
    from backend.services.analysis_service import analyze_technical
    return analyze_technical(ticker)


def _load_news(ticker: str) -> dict | None:
    from backend.services.news_service import analyze_news
    return analyze_news(ticker)


def _latest_history(ticker: str, kind: str) -> dict | None:
    from backend.infra.schema import AnalysisResult
    with get_session() as s:
        row = (s.query(AnalysisResult).filter_by(ticker=ticker, kind=kind)
               .order_by(AnalysisResult.id.desc()).first())
    return json.loads(row.content_json) if row else None


def generate_stock_report(ticker: str) -> dict:
    """종목분석 리포트 docx 생성 (FR-04-41) — reports/YYYY-MM/ 저장 + 등록."""
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    today = date.today()

    # 표지
    doc.add_heading(f"종목분석 리포트 — {ticker}", 0)
    doc.add_paragraph(f"생성일: {today.isoformat()} · 분석 기준일: T-1")
    doc.add_paragraph(_DISCLAIMER)

    # 종합 판단 (이력)
    judgment = _latest_history(ticker, "comprehensive")
    if judgment:
        doc.add_heading("1. 종합 판단", 1)
        doc.add_paragraph(f"제안: {judgment.get('recommendation') or '(서술형 참조)'}")
        if judgment.get("fair_value_current"):
            doc.add_paragraph(f"적정가(현재): {judgment['fair_value_current']:,} / "
                              f"적정가(12개월): {judgment.get('fair_value_future') or '-'}")
        if judgment.get("plan"):
            doc.add_paragraph(f"투자 방안: {judgment['plan']}")
        if judgment.get("rationale"):
            doc.add_paragraph(f"근거: {judgment['rationale']}")
        if judgment.get("assumptions"):
            doc.add_paragraph("가정: " + " / ".join(judgment["assumptions"]))
        if judgment.get("narrative"):
            doc.add_paragraph(judgment["narrative"])

    # 분석 A
    try:
        fund = _load_fundamental(ticker)
    except Exception as e:
        fund = None
        doc.add_paragraph(f"(재무 분석 조회 실패: {e})")
    if fund:
        doc.add_heading("2. 기초(재무) 분석", 1)
        doc.add_paragraph(f"Tier 1 판정: {fund['evaluation']['tier1']['verdict']}")
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(("연도", "매출액", "영업이익", "순이익")):
            t.rows[0].cells[i].text = h
        for f in fund["financials"]:
            row = t.add_row().cells
            row[0].text = str(f["year"])
            for i, k in enumerate(("revenue", "operating_profit", "net_income"), 1):
                row[i].text = f"{f.get(k):,}" if f.get(k) is not None else "-"
        doc.add_paragraph("")
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = "Light Grid Accent 1"
        for i, h in enumerate(("지표", "값", "기준", "판정")):
            t2.rows[0].cells[i].text = h
        for item in fund["evaluation"]["items"]:
            row = t2.add_row().cells
            row[0].text = item["label"]
            row[1].text = str(item["value"]) if item["value"] is not None else "-"
            row[2].text = f"{'≥' if item['direction'] == 'min' else '≤'} {item['threshold']}"
            row[3].text = item["status"]

    # 분석 B
    try:
        tech = _load_technical(ticker)
    except Exception:
        tech = None
    if tech:
        doc.add_heading("3. 기술적 분석", 1)
        doc.add_paragraph(f"시그널: {tech['signal']['verdict']} "
                          f"(배열 {tech.get('ma_alignment') or '-'}, RSI {tech.get('rsi') or '-'})")
        for reason in tech["signal"]["reasons"]:
            doc.add_paragraph(f"- {reason}")

    # 분석 C + 출처
    try:
        news = _load_news(ticker)
    except Exception:
        news = None
    if news:
        doc.add_heading("4. 뉴스·공시", 1)
        for d in news.get("disclosures", [])[:5]:
            doc.add_paragraph(f"[공시 {d['date']}] {d['title']} — {d['link']}")
        for n in news.get("news", [])[:8]:
            tag = f"[{n.get('sentiment')}] " if n.get("sentiment") else ""
            doc.add_paragraph(f"{tag}{n['title']} ({n['source']}, {n['date']}) — {n['link']}")

    # 토론 (이력)
    deb = _latest_history(ticker, "debate")
    if deb:
        doc.add_heading("5. AI 토론", 1)
        doc.add_paragraph(f"[강세론] {deb['bull']}")
        doc.add_paragraph(f"[약세론] {deb['bear']}")
        doc.add_paragraph(f"[결론] {deb['conclusion']}")

    deep = _latest_history(ticker, "deep")
    if deep:
        doc.add_heading("6. 딥리서치", 1)
        doc.add_paragraph(deep["content"])

    doc.add_paragraph("")
    p = doc.add_paragraph(_DISCLAIMER)
    p.runs[0].font.size = Pt(9)

    # 저장 + 등록 (FR-09-03)
    month_dir = REPORT_DIR / today.strftime("%Y-%m")
    month_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{ticker}_{today.isoformat()}_{datetime.now().strftime('%H%M%S')}.docx"
    path = month_dir / filename
    doc.save(path)
    relpath = f"{today.strftime('%Y-%m')}/{filename}"
    with get_session() as s:
        rep = Report(ticker=ticker, kind="stock", filename=filename,
                     relpath=relpath, created_at=datetime.now())
        s.add(rep)
        s.commit()
        rid = rep.id
    logger.info("리포트 생성: %s", path)
    return {"id": rid, "ticker": ticker, "relpath": relpath}


def list_reports() -> list[dict]:
    with get_session() as s:
        rows = s.query(Report).order_by(Report.id.desc()).limit(50).all()
    return [{"id": r.id, "ticker": r.ticker, "kind": r.kind, "filename": r.filename,
             "created_at": r.created_at.isoformat(timespec="seconds")} for r in rows]


def report_path(report_id: int) -> Path | None:
    with get_session() as s:
        r = s.get(Report, report_id)
    if not r:
        return None
    p = REPORT_DIR / r.relpath
    return p if p.exists() else None


def generate_rebalance_report() -> dict:
    """리밸런싱 리포트 (FR-05-17) — 최근 제안 이력 기반."""
    import docx
    from docx.shared import Pt
    from datetime import date, datetime

    proposal = _latest_history("PORTFOLIO", "rebalance")
    if not proposal:
        raise ValueError("리밸런싱 제안 이력이 없습니다 — 먼저 제안을 실행하세요")

    doc = docx.Document()
    today = date.today()
    doc.add_heading("자산 리밸런싱 리포트", 0)
    doc.add_paragraph(f"생성일: {today.isoformat()}")
    doc.add_paragraph(_DISCLAIMER)

    doc.add_heading("1. 목표 대비 이탈", 1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(("항목", "현재", "목표", "이탈")):
        t.rows[0].cells[i].text = h
    for d in proposal.get("deviations", []):
        row = t.add_row().cells
        row[0].text = d["label"]
        row[1].text = f"{d['current_pct']}%"
        row[2].text = f"{d['target_pct']}%"
        row[3].text = f"{d['deviation_pp']:+}%p"

    doc.add_heading("2. 매매 제안", 1)
    if proposal.get("summary"):
        doc.add_paragraph(proposal["summary"])
    t2 = doc.add_table(rows=1, cols=5)
    t2.style = "Light Grid Accent 1"
    for i, h in enumerate(("종목", "액션", "수량", "예상금액", "근거")):
        t2.rows[0].cells[i].text = h
    for a in proposal.get("actions", []):
        row = t2.add_row().cells
        row[0].text = f"{a.get('name')}({a.get('ticker')})"
        row[1].text = a.get("action", "")
        row[2].text = f"{a.get('qty', 0):,}"
        row[3].text = f"{a.get('est_amount', 0):,}"
        row[4].text = a.get("rationale", "")
    if proposal.get("narrative"):
        doc.add_paragraph(proposal["narrative"])

    doc.add_heading("3. 실행 전/후 비중", 1)
    for r in proposal.get("before_after", []):
        doc.add_paragraph(f"{r['label']}: {r['before_pct']}% → {r['after_pct']}%")
    for w in proposal.get("warnings", []):
        doc.add_paragraph(f"⚠ 검증 경고: {w}")

    p2 = doc.add_paragraph(_DISCLAIMER)
    p2.runs[0].font.size = Pt(9)

    month_dir = REPORT_DIR / today.strftime("%Y-%m")
    month_dir.mkdir(parents=True, exist_ok=True)
    filename = f"rebalance_{today.isoformat()}_{datetime.now().strftime('%H%M%S')}.docx"
    doc.save(month_dir / filename)
    relpath = f"{today.strftime('%Y-%m')}/{filename}"
    with get_session() as s:
        rep = Report(ticker="PORTFOLIO", kind="rebalance", filename=filename,
                     relpath=relpath, created_at=datetime.now())
        s.add(rep)
        s.commit()
        rid = rep.id
    return {"id": rid, "relpath": relpath}
